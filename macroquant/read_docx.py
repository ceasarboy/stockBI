from docx import Document

doc = Document(r'f:\stock\macroquant\docs\项目需求讨论.docx')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

for table in doc.tables:
    print("\n--- TABLE ---")
    for row in table.rows:
        row_text = [cell.text for cell in row.cells]
        print(" | ".join(row_text))
